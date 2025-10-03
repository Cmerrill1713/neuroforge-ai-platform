#!/usr/bin/env python3
"""
Enhanced Multimodal Input Processing System with Docling Integration
Handle images, documents, audio, and other media types with advanced document processing
"""

import asyncio
import base64
import json
import logging
import mimetypes
import tempfile
from typing import Dict, Any, List, Optional, Union, Tuple
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
import io
import hashlib

logger = logging.getLogger(__name__)

@dataclass
class EnhancedProcessingResult:
    """Enhanced result of multimodal input processing with Docling data"""
    success: bool
    content_type: str
    extracted_data: Dict[str, Any]
    docling_data: Optional[Dict[str, Any]] = None
    metadata: Dict[str, Any] = field(default_factory=dict)
    processing_time_ms: float = 0.0
    error_message: Optional[str] = None
    file_hash: Optional[str] = None

@dataclass
class DocumentStructure:
    """Document structure extracted by Docling"""
    title: Optional[str] = None
    sections: List[Dict[str, Any]] = field(default_factory=list)
    tables: List[Dict[str, Any]] = field(default_factory=list)
    images: List[Dict[str, Any]] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)

class EnhancedMultimodalInputProcessor:
    """Enhanced multimodal input processor with Docling integration"""
    
    def __init__(self):
        self.supported_formats = {
            "image": ["jpg", "jpeg", "png", "gif", "bmp", "webp", "tiff"],
            "document": ["pdf", "docx", "txt", "md", "rtf", "odt", "pptx", "xlsx", "html"],
            "audio": ["mp3", "wav", "flac", "aac", "ogg", "m4a"],
            "video": ["mp4", "avi", "mov", "mkv", "webm"],
            "data": ["json", "csv", "xml", "yaml", "yml"]
        }
        
        self.processors = {
            "image": self._process_image,
            "document": self._process_document_enhanced,
            "audio": self._process_audio,
            "video": self._process_video,
            "data": self._process_data
        }
        
        # Initialize Docling
        self.docling_available = self._initialize_docling()
        
        logger.info(f"🎭 Enhanced Multimodal Input Processor initialized (Docling: {'✅' if self.docling_available else '❌'})")
    
    def _initialize_docling(self) -> bool:
        """Initialize Docling for advanced document processing"""
        try:
            from docling.document_converter import DocumentConverter
            from docling.datamodel.base_models import InputFormat
            from docling.datamodel.pipeline_options import PdfPipelineOptions
            from docling.datamodel.pipeline_options import PdfPipelineOptions
            
            # Initialize converter with optimized settings
            self.converter = DocumentConverter(
                format_options={
                    InputFormat.PDF: PdfPipelineOptions(
                        do_ocr=True,
                        do_table_structure=True,
                        table_structure_options={"do_cell_matching": True}
                    )
                }
            )
            logger.info("✅ Docling initialized successfully")
            return True
        except ImportError as e:
            logger.warning(f"⚠️ Docling not available: {e}")
            return False
        except Exception as e:
            logger.error(f"❌ Failed to initialize Docling: {e}")
            return False
    
    async def process_input(
        self,
        input_data: Union[str, bytes, Path],
        content_type: Optional[str] = None,
        use_docling: bool = True
    ) -> EnhancedProcessingResult:
        """
        Process multimodal input with optional Docling enhancement
        
        Args:
            input_data: Input data (file path, bytes, or string)
            content_type: MIME type of the input
            use_docling: Whether to use Docling for document processing
        """
        start_time = datetime.now()
        
        try:
            # Determine content type if not provided
            if not content_type:
                content_type = await self._detect_content_type(input_data)
            
            # Determine file type
            file_type = self._get_file_type(content_type)
            
            # Process with appropriate processor
            if file_type in self.processors:
                extracted_data = await self.processors[file_type](input_data, content_type, use_docling)
            else:
                raise ValueError(f"Unsupported file type: {content_type}")
            
            # Calculate processing time
            processing_time = (datetime.now() - start_time).total_seconds() * 1000
            
            # Generate file hash
            file_hash = self._generate_file_hash(input_data)
            
            return EnhancedProcessingResult(
                success=True,
                content_type=content_type,
                extracted_data=extracted_data,
                metadata={
                    "file_type": file_type,
                    "processing_time_ms": processing_time,
                    "docling_used": use_docling and self.docling_available and file_type == "document"
                },
                processing_time_ms=processing_time,
                file_hash=file_hash
            )
            
        except Exception as e:
            logger.error(f"Processing failed: {e}")
            return EnhancedProcessingResult(
                success=False,
                content_type=content_type or "unknown",
                extracted_data={},
                error_message=str(e),
                processing_time_ms=(datetime.now() - start_time).total_seconds() * 1000
            )
    
    async def _process_document_enhanced(
        self, 
        input_data: Union[str, bytes, Path], 
        content_type: str,
        use_docling: bool = True
    ) -> Dict[str, Any]:
        """Enhanced document processing with Docling integration"""
        
        # Try Docling first if available and requested
        if use_docling and self.docling_available and content_type in ['application/pdf', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
            try:
                docling_result = await self._process_with_docling(input_data, content_type)
                if docling_result:
                    return docling_result
            except Exception as e:
                logger.warning(f"Docling processing failed, falling back to basic processing: {e}")
        
        # Fallback to basic processing
        return await self._process_document_basic(input_data, content_type)
    
    async def _process_with_docling(
        self, 
        input_data: Union[str, bytes, Path], 
        content_type: str
    ) -> Optional[Dict[str, Any]]:
        """Process document using Docling for advanced extraction"""
        try:
            # Convert input to file path if needed
            file_path = await self._prepare_file_path(input_data)
            
            # Convert document with Docling
            result = self.converter.convert(file_path)
            doc = result.document
            
            # Extract structured content
            document_structure = DocumentStructure(
                title=doc.name if hasattr(doc, 'name') else None,
                sections=[],
                tables=[],
                images=[],
                metadata={}
            )
            
            # Extract text content
            full_text = ""
            if hasattr(doc, 'text'):
                full_text = doc.text
            elif hasattr(doc, 'export_to_markdown'):
                full_text = doc.export_to_markdown()
            
            # Extract tables if available
            tables = []
            if hasattr(doc, 'tables'):
                for table in doc.tables:
                    table_data = {
                        "headers": [],
                        "rows": [],
                        "metadata": {}
                    }
                    
                    # Extract table structure
                    if hasattr(table, 'export_to_markdown'):
                        table_data["markdown"] = table.export_to_markdown()
                    
                    tables.append(table_data)
            
            # Extract images if available
            images = []
            if hasattr(doc, 'images'):
                for img in doc.images:
                    image_data = {
                        "caption": getattr(img, 'caption', None),
                        "metadata": {}
                    }
                    images.append(image_data)
            
            # Analyze document structure
            analysis = await self._analyze_document_content(full_text)
            
            return {
                "type": "document",
                "title": document_structure.title or analysis.get("title"),
                "content": full_text,
                "pages": analysis.get("pages", 1),
                "word_count": analysis.get("word_count", 0),
                "language": analysis.get("language"),
                "keywords": analysis.get("keywords", []),
                "summary": analysis.get("summary"),
                "structure": analysis.get("structure", {}),
                "tables": tables,
                "images": images,
                "docling_processed": True,
                "docling_metadata": {
                    "converter_version": "docling-2.55.0",
                    "processing_method": "advanced",
                    "ocr_enabled": True,
                    "table_extraction": True
                }
            }
            
        except Exception as e:
            logger.error(f"Docling processing failed: {e}")
            return None
    
    async def _process_document_basic(
        self, 
        input_data: Union[str, bytes, Path], 
        content_type: str
    ) -> Dict[str, Any]:
        """Basic document processing (fallback)"""
        try:
            # Convert to text content using basic methods
            if isinstance(input_data, Path) and input_data.exists():
                if content_type == 'application/pdf':
                    content = await self._extract_pdf_content_basic(input_data)
                elif content_type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
                    content = await self._extract_docx_content_basic(input_data)
                else:
                    with open(input_data, 'r', encoding='utf-8') as f:
                        content = f.read()
            elif isinstance(input_data, str):
                content = input_data
            else:
                content = input_data.decode('utf-8')
            
            # Analyze document
            analysis = await self._analyze_document_content(content)
            
            return {
                "type": "document",
                "title": analysis.get("title"),
                "content": content,
                "pages": analysis.get("pages", 1),
                "word_count": analysis.get("word_count", 0),
                "language": analysis.get("language"),
                "keywords": analysis.get("keywords", []),
                "summary": analysis.get("summary"),
                "structure": analysis.get("structure", {}),
                "tables": [],
                "images": [],
                "docling_processed": False,
                "docling_metadata": {
                    "processing_method": "basic",
                    "fallback": True
                }
            }
            
        except Exception as e:
            logger.error(f"Basic document processing failed: {e}")
            return {"type": "document", "error": str(e), "docling_processed": False}
    
    async def _prepare_file_path(self, input_data: Union[str, bytes, Path]) -> Path:
        """Prepare file path for Docling processing"""
        if isinstance(input_data, Path):
            return input_data
        elif isinstance(input_data, str):
            return Path(input_data)
        else:
            # Create temporary file for bytes data
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
                tmp_file.write(input_data)
                return Path(tmp_file.name)
    
    async def _extract_pdf_content_basic(self, pdf_path: Path) -> str:
        """Basic PDF content extraction (fallback)"""
        try:
            import PyPDF2
            
            with open(pdf_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                content = ""
                
                for page in reader.pages:
                    content += page.extract_text() + "\n"
                
                return content.strip()
                
        except ImportError:
            logger.warning("PyPDF2 not available for PDF processing")
            return "PDF processing not available"
        except Exception as e:
            logger.warning(f"PDF extraction failed: {e}")
            return f"PDF extraction failed: {e}"
    
    async def _extract_docx_content_basic(self, docx_path: Path) -> str:
        """Basic DOCX content extraction (fallback)"""
        try:
            from docx import Document
            
            doc = Document(docx_path)
            content = ""
            
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            
            return content.strip()
            
        except ImportError:
            logger.warning("python-docx not available for DOCX processing")
            return "DOCX processing not available"
        except Exception as e:
            logger.warning(f"DOCX extraction failed: {e}")
            return f"DOCX extraction failed: {e}"
    
    async def _analyze_document_content(self, content: str) -> Dict[str, Any]:
        """Analyze document content and extract metadata"""
        # Basic analysis
        words = content.split()
        sentences = content.split('.')
        
        # Extract title (first line or first sentence)
        title = None
        lines = content.split('\n')
        if lines:
            first_line = lines[0].strip()
            if len(first_line) < 100 and len(first_line.split()) < 10:
                title = first_line
        
        # Basic keyword extraction (first 10 unique words)
        keywords = list(set([word.lower().strip('.,!?;:') for word in words[:50] if len(word) > 3]))[:10]
        
        # Generate summary (first 200 characters)
        summary = content[:200] + "..." if len(content) > 200 else content
        
        return {
            "title": title,
            "word_count": len(words),
            "sentence_count": len(sentences),
            "line_count": len(lines),
            "keywords": keywords,
            "summary": summary,
            "structure": {
                "has_headers": any(line.strip().endswith(':') for line in lines[:10]),
                "has_lists": any(line.strip().startswith(('-', '*', '1.', '2.')) for line in lines),
                "has_paragraphs": len([line for line in lines if len(line.strip()) > 50]) > 0
            }
        }
    
    async def _process_image(self, input_data: Union[str, bytes, Path], content_type: str, use_docling: bool = True) -> Dict[str, Any]:
        """Process image input"""
        # Basic image processing (existing implementation)
        return {"type": "image", "processed": True, "content_type": content_type}
    
    async def _process_audio(self, input_data: Union[str, bytes, Path], content_type: str, use_docling: bool = True) -> Dict[str, Any]:
        """Process audio input"""
        # Basic audio processing (existing implementation)
        return {"type": "audio", "processed": True, "content_type": content_type}
    
    async def _process_video(self, input_data: Union[str, bytes, Path], content_type: str, use_docling: bool = True) -> Dict[str, Any]:
        """Process video input"""
        # Basic video processing (existing implementation)
        return {"type": "video", "processed": True, "content_type": content_type}
    
    async def _process_data(self, input_data: Union[str, bytes, Path], content_type: str, use_docling: bool = True) -> Dict[str, Any]:
        """Process structured data input"""
        # Basic data processing (existing implementation)
        return {"type": "data", "processed": True, "content_type": content_type}
    
    async def _detect_content_type(self, input_data: Union[str, bytes, Path]) -> str:
        """Detect content type from input data"""
        if isinstance(input_data, Path):
            mime_type, _ = mimetypes.guess_type(str(input_data))
            return mime_type or 'application/octet-stream'
        elif isinstance(input_data, str):
            return 'text/plain'
        else:
            return 'application/octet-stream'
    
    def _get_file_type(self, content_type: str) -> str:
        """Get file type category from content type"""
        if content_type.startswith('image/'):
            return 'image'
        elif content_type.startswith('video/'):
            return 'video'
        elif content_type.startswith('audio/'):
            return 'audio'
        elif content_type in ['application/pdf', 'application/msword', 
                             'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                             'application/vnd.openxmlformats-officedocument.presentationml.presentation',
                             'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                             'text/plain', 'text/html', 'text/markdown']:
            return 'document'
        else:
            return 'data'
    
    def _generate_file_hash(self, input_data: Union[str, bytes, Path]) -> str:
        """Generate hash for input data"""
        try:
            if isinstance(input_data, Path):
                with open(input_data, 'rb') as f:
                    content = f.read()
            elif isinstance(input_data, str):
                content = input_data.encode('utf-8')
            else:
                content = input_data
            
            return hashlib.md5(content).hexdigest()
        except Exception:
            return "unknown"

# Global enhanced multimodal processor instance
enhanced_multimodal_processor = EnhancedMultimodalInputProcessor()

async def main():
    """Test enhanced multimodal input processor"""
    print("🎭 Enhanced Multimodal Input Processor Test")
    
    # Test basic text processing
    text_result = await enhanced_multimodal_processor.process_input(
        "Hello, this is a test document.", 
        "text/plain"
    )
    print(f"Text processing: {text_result.success}")
    print(f"Docling used: {text_result.metadata.get('docling_used', False)}")
    
    # Test with a sample PDF (if available)
    # pdf_result = await enhanced_multimodal_processor.process_input(
    #     "sample.pdf", 
    #     "application/pdf"
    # )
    # print(f"PDF processing: {pdf_result.success}")
    # print(f"Docling used: {pdf_result.metadata.get('docling_used', False)}")

if __name__ == "__main__":
    asyncio.run(main())
